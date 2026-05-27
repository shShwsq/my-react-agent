import os
import re
import logging
from typing import Optional
from pathlib import Path
from dataclasses import dataclass

logger = logging.getLogger(__name__)

FILES_DIR = Path("storage/agent_files")

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("[FileOperations] python-docx not installed, Word file support disabled")

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("[FileOperations] pdfplumber not installed, PDF text extraction disabled")

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    logger.warning("[FileOperations] openpyxl not installed, Excel file support disabled")


@dataclass
class FileResult:
    success: bool
    filename: Optional[str] = None
    content: Optional[str] = None
    file_path: Optional[str] = None
    relative_path: Optional[str] = None
    file_size: int = 0
    error: Optional[str] = None
    message: Optional[str] = None


def ensure_storage_dir():
    if not FILES_DIR.exists():
        FILES_DIR.mkdir(parents=True, exist_ok=True)


async def write_file(
    filename: str,
    content: str,
    room_id: str,
    user_id: int,
    mode: str = "write"
) -> FileResult:
    try:
        ensure_storage_dir()
        
        outputs_dir = FILES_DIR / str(user_id) / room_id / "outputs"
        outputs_dir.mkdir(parents=True, exist_ok=True)
        
        file_path = outputs_dir / filename
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        if ext == 'docx':
            return await _write_docx(file_path, filename, content, mode)
        else:
            return await _write_text(file_path, filename, content, mode)
        
    except Exception as e:
        logger.error(f"[FileOperations] Write error: {e}")
        return FileResult(
            success=False,
            error=str(e)
        )


async def _write_text(file_path: Path, filename: str, content: str, mode: str) -> FileResult:
    write_mode = "w" if mode == "write" else "a"
    with open(file_path, write_mode, encoding="utf-8") as f:
        f.write(content)
    
    file_size = file_path.stat().st_size
    
    return FileResult(
        success=True,
        filename=filename,
        file_path=str(file_path),
        relative_path=file_path.relative_to(FILES_DIR).as_posix(),
        file_size=file_size,
        message=f"文件 {filename} 已成功保存"
    )


async def _write_docx(file_path: Path, filename: str, content: str, mode: str) -> FileResult:
    if not DOCX_AVAILABLE:
        return FileResult(
            success=False,
            error="Word 文件支持未安装，请运行: pip install python-docx"
        )
    
    if mode == "append" and file_path.exists():
        doc = Document(str(file_path))
        _add_markdown_content(doc, content)
    else:
        doc = Document()
        _add_markdown_content(doc, content)
    
    doc.save(str(file_path))
    file_size = file_path.stat().st_size
    
    return FileResult(
        success=True,
        filename=filename,
        file_path=str(file_path),
        relative_path=file_path.relative_to(FILES_DIR).as_posix(),
        file_size=file_size,
        message=f"Word 文件 {filename} 已成功保存"
    )


def _add_markdown_content(doc, content: str):
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        if re.match(r'^#{1}\s+', line):
            text = re.sub(r'^#\s+', '', line)
            para = doc.add_heading(text, level=1)
        elif re.match(r'^#{2}\s+', line):
            text = re.sub(r'^##\s+', '', line)
            para = doc.add_heading(text, level=2)
        elif re.match(r'^#{3}\s+', line):
            text = re.sub(r'^###\s+', '', line)
            para = doc.add_heading(text, level=3)
        elif re.match(r'^#{4,6}\s+', line):
            text = re.sub(r'^#{4,6}\s+', '', line)
            para = doc.add_heading(text, level=4)
        elif line.strip() == '':
            doc.add_paragraph('')
        else:
            para = doc.add_paragraph()
            _add_formatted_text(para, line)
        
        i += 1


def _add_formatted_text(paragraph, text: str):
    parts = []
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|[^*`]+)'
    matches = re.findall(pattern, text)
    
    for match in matches:
        if match.startswith('**') and match.endswith('**'):
            run = paragraph.add_run(match[2:-2])
            run.bold = True
        elif match.startswith('*') and match.endswith('*') and not match.startswith('**'):
            run = paragraph.add_run(match[1:-1])
            run.italic = True
        elif match.startswith('`') and match.endswith('`'):
            run = paragraph.add_run(match[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(match)


async def read_file(
    filename: str,
    room_id: str,
    user_id: int,
    folder: str = "outputs"
) -> FileResult:
    try:
        folder_dir = FILES_DIR / str(user_id) / room_id / folder
        
        # 如果指定文件夹存在且文件也在其中，直接用
        if folder_dir.exists():
            file_path = folder_dir / filename
            if file_path.exists():
                pass  # 文件找到了
            else:
                file_path = None
        else:
            file_path = None
        
        # 如果没找到，自动搜索 outputs 和 uploads
        if file_path is None:
            searched_folders = ["outputs", "uploads"]
            for search_folder in searched_folders:
                search_path = FILES_DIR / str(user_id) / room_id / search_folder / filename
                if search_path.exists():
                    file_path = search_path
                    folder = search_folder
                    break
            else:
                return FileResult(
                    success=False,
                    error=f"文件 {filename} 不存在于 outputs 或 uploads 文件夹中"
                )
        
        ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
        
        if ext == 'docx':
            content = _extract_docx(file_path)
        elif ext == 'pdf':
            content = _extract_pdf(file_path)
        elif ext == 'xlsx':
            content = _extract_xlsx(file_path)
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        
        return FileResult(
            success=True,
            filename=filename,
            content=content,
            file_size=file_path.stat().st_size
        )
        
    except Exception as e:
        logger.error(f"[FileOperations] Read error: {e}")
        return FileResult(
            success=False,
            error=str(e)
        )


def _extract_docx(file_path: Path) -> str:
    """从 docx 文件提取文本"""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx 未安装，无法读取 Word 文件，请运行: pip install python-docx")
    
    doc = Document(str(file_path))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    
    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                lines.append(row_text)
    
    return "\n".join(lines)


def _extract_pdf(file_path: Path) -> str:
    """从 pdf 文件提取文本"""
    if not PDF_AVAILABLE:
        raise RuntimeError("pdfplumber 未安装，无法读取 PDF 文件，请运行: pip install pdfplumber")
    
    lines = []
    with pdfplumber.open(str(file_path)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                if len(pdf.pages) > 1:
                    lines.append(f"--- 第 {i + 1} 页 ---")
                lines.append(text)
            
            # 提取表格
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    row_text = " | ".join(str(cell or "") for cell in row)
                    if row_text.strip(" |"):
                        lines.append(row_text)
    
    return "\n".join(lines)


def _extract_xlsx(file_path: Path) -> str:
    """从 xlsx 文件提取文本"""
    if not XLSX_AVAILABLE:
        raise RuntimeError("openpyxl 未安装，无法读取 Excel 文件，请运行: pip install openpyxl")
    
    wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
    lines = []
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        if len(wb.sheetnames) > 1:
            lines.append(f"=== 工作表: {sheet_name} ===")
        
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip(" |"):
                lines.append(row_text)
    
    wb.close()
    return "\n".join(lines)


def list_files(room_id: str, user_id: int, folder: str = "all") -> list[dict]:
    try:
        room_dir = FILES_DIR / str(user_id) / room_id
        
        if not room_dir.exists():
            return []
        
        files = []
        
        if folder == "all":
            folders_to_search = ["outputs", "uploads"]
        else:
            folders_to_search = [folder]
        
        for folder_name in folders_to_search:
            folder_dir = room_dir / folder_name
            if folder_dir.exists():
                for file_path in folder_dir.iterdir():
                    if file_path.is_file():
                        files.append({
                            "filename": file_path.name,
                            "folder": folder_name,
                            "file_size": file_path.stat().st_size,
                            "relative_path": f"{user_id}/{room_id}/{folder_name}/{file_path.name}",
                            "created_at": file_path.stat().st_ctime
                        })
        
        return files
        
    except Exception as e:
        logger.error(f"[FileOperations] List error: {e}")
        return []


def get_file_path(filename: str, room_id: str, user_id: int, folder: str = "outputs") -> Optional[Path]:
    file_path = FILES_DIR / str(user_id) / room_id / folder / filename
    return file_path if file_path.exists() else None


from app.tasks.file_operations.registry import (
    BaseFileOperation,
    FileOperationResult,
    file_operation_registry
)


class WriteFileOperation(BaseFileOperation):
    name = "file_write"
    description = "写入文件。支持 txt、md、json、csv、html、py、js 等文本文件，以及 docx（Word文档，支持 Markdown 语法转换）"
    
    def get_parameters(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "文件名，包含扩展名。如 'report.md'、'data.json'、'document.docx'"
                },
                "content": {
                    "type": "string",
                    "description": "要写入的文件内容。docx 文件支持 Markdown 语法：# 标题、**粗体**、*斜体*、`代码`"
                },
                "mode": {
                    "type": "string",
                    "enum": ["write", "append"],
                    "description": "写入模式：'write' 覆盖写入（默认），'append' 追加写入"
                }
            },
            "required": ["filename", "content"]
        }
    
    async def execute(self, **kwargs) -> FileOperationResult:
        filename = kwargs.get("filename", "")
        content = kwargs.get("content", "")
        mode = kwargs.get("mode", "write")
        room_id = kwargs.get("room_id", "default")
        user_id = kwargs.get("user_id", 0)
        
        result = await write_file(
            filename=filename,
            content=content,
            room_id=room_id,
            user_id=user_id,
            mode=mode
        )
        
        return FileOperationResult(
            success=result.success,
            output={
                "filename": result.filename,
                "file_path": result.file_path,
                "relative_path": result.relative_path,
                "file_size": result.file_size,
                "message": result.message
            },
            error=result.error
        )


class ReadFileOperation(BaseFileOperation):
    name = "file_read"
    description = "读取文件内容。支持多种文件格式：txt、md、json、csv、html、py、js 等文本文件，以及 PDF（提取文字和表格）、DOCX（Word文档）、XLSX（Excel表格）。会自动在 outputs 和 uploads 文件夹中搜索文件"
    
    def get_parameters(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "要读取的文件名"
                },
                "folder": {
                    "type": "string",
                    "enum": ["outputs", "uploads"],
                    "description": "文件夹名称，默认 outputs。如果不指定，会自动在 outputs 和 uploads 中搜索"
                }
            },
            "required": ["filename"]
        }
    
    async def execute(self, **kwargs) -> FileOperationResult:
        filename = kwargs.get("filename", "")
        folder = kwargs.get("folder", "outputs")
        room_id = kwargs.get("room_id", "default")
        user_id = kwargs.get("user_id", 0)
        
        result = await read_file(
            filename=filename,
            room_id=room_id,
            user_id=user_id,
            folder=folder
        )
        
        return FileOperationResult(
            success=result.success,
            output={
                "filename": result.filename,
                "content": result.content,
                "file_size": result.file_size,
                "message": f"成功读取文件 {result.filename}" if result.success else None
            },
            error=result.error
        )


class ListFilesOperation(BaseFileOperation):
    name = "file_list"
    description = "列出当前房间内的文件，可指定 outputs 或 uploads 文件夹"
    
    def get_parameters(self) -> dict:
        return {
            "type": "object",
            "properties": {
                "folder": {
                    "type": "string",
                    "enum": ["all", "outputs", "uploads"],
                    "description": "文件夹名称：'all' 列出所有文件（默认），'outputs' 只列出智能体生成的文件，'uploads' 只列出上传的文件"
                }
            },
            "required": []
        }
    
    async def execute(self, **kwargs) -> FileOperationResult:
        folder = kwargs.get("folder", "all")
        room_id = kwargs.get("room_id", "default")
        user_id = kwargs.get("user_id", 0)
        
        try:
            files = list_files(room_id=room_id, user_id=user_id, folder=folder)
            
            return FileOperationResult(
                success=True,
                output={
                    "folder": folder,
                    "count": len(files),
                    "files": files,
                    "message": f"找到 {len(files)} 个文件" if files else "文件夹为空"
                }
            )
        except Exception as e:
            logger.error(f"[ListFilesOperation] Error: {e}")
            return FileOperationResult(
                success=False,
                output=None,
                error=str(e)
            )


file_operation_registry.register(WriteFileOperation())
file_operation_registry.register(ReadFileOperation())
file_operation_registry.register(ListFilesOperation())
